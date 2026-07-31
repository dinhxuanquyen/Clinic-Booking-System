from __future__ import annotations

import json
import sys
import tempfile
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from lxml import etree


WORD_MACRO_MAIN = "application/vnd.ms-word.document.macroEnabled.main+xml"
WORD_DOCX_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"


def make_macro_free_copy(source: Path, destination: Path) -> Path:
    with zipfile.ZipFile(source, "r") as source_archive:
        with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as destination_archive:
            for info in source_archive.infolist():
                if info.filename == "word/vbaProject.bin":
                    continue
                data = source_archive.read(info.filename)
                if info.filename == "[Content_Types].xml":
                    data = data.replace(WORD_MACRO_MAIN.encode(), WORD_DOCX_MAIN.encode())
                    xml = etree.fromstring(data)
                    for override in list(xml):
                        if override.get("PartName") == "/word/vbaProject.bin":
                            xml.remove(override)
                    data = etree.tostring(xml, xml_declaration=True, encoding="UTF-8", standalone=True)
                elif info.filename == "word/_rels/document.xml.rels":
                    xml = etree.fromstring(data)
                    for relationship in list(xml):
                        if relationship.get("Type", "").endswith("/vbaProject"):
                            xml.remove(relationship)
                    data = etree.tostring(xml, xml_declaration=True, encoding="UTF-8", standalone=True)
                destination_archive.writestr(info, data)
    return destination


def main() -> None:
    path = Path(sys.argv[1]).resolve()
    out_path = Path(sys.argv[2]).resolve()
    normalized_path = Path(sys.argv[3]).resolve() if len(sys.argv) > 3 else None
    if path.suffix.lower() == ".docm":
        if normalized_path is None:
            normalized_path = Path(tempfile.gettempdir()) / f"{path.stem}-inspection.docx"
        make_macro_free_copy(path, normalized_path)
        document_path = normalized_path
    else:
        document_path = path
    doc = Document(document_path)

    paragraphs = []
    style_counts: Counter[str] = Counter()
    headings = []
    nonempty_index = 0
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        style_counts[style] += 1
        if text:
            nonempty_index += 1
            paragraphs.append(
                {
                    "index": index,
                    "nonempty_index": nonempty_index,
                    "style": style,
                    "text": text,
                }
            )
            lowered_style = style.lower()
            if "heading" in lowered_style or "title" in lowered_style or "tiêu đề" in lowered_style:
                headings.append(
                    {
                        "index": index,
                        "style": style,
                        "text": text,
                    }
                )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(
            {
                "index": table_index,
                "rows": len(table.rows),
                "columns": max((len(row.cells) for row in table.rows), default=0),
                "preview": rows[:5],
            }
        )

    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        comments = [name for name in names if name.startswith("word/comments")]
        footnotes = "word/footnotes.xml" in names
        endnotes = "word/endnotes.xml" in names
        has_vba = "word/vbaProject.bin" in names
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        tracked_insertions = len(document_xml.xpath(".//w:ins", namespaces=namespace))
        tracked_deletions = len(document_xml.xpath(".//w:del", namespaces=namespace))

    result = {
        "path": str(path),
        "normalized_path": str(normalized_path) if normalized_path else None,
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "media_count": len(media),
        "media_files": media,
        "has_vba": has_vba,
        "comments_parts": comments,
        "has_footnotes": footnotes,
        "has_endnotes": endnotes,
        "tracked_insertions": tracked_insertions,
        "tracked_deletions": tracked_deletions,
        "style_counts": style_counts.most_common(),
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables,
    }
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({key: result[key] for key in result if key not in {"paragraphs", "tables", "media_files"}}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
