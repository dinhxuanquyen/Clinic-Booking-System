from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

OUT_PATH = r"C:\Users\DELL\Documents\Codex\2026-05-27\clinic-booking\4.3_thiet_ke_api_v2.docx"


def set_font(run, name="Times New Roman", size=13, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_paragraph(p, align=None, space_after=6, line_spacing=1.5):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=13)
    style_paragraph(p)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    r = p.add_run(text)
    set_font(r, size=14 if level == 1 else 13, bold=True)
    style_paragraph(p, space_after=4)
    return p


def set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = widths[i]
        set_cell_shading(c, "F2F2F2")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = widths[i]
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_font(r, size=11)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
    return table


doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("4.3. Thiết kế API")
set_font(r, size=16, bold=True)
style_paragraph(p, space_after=10)

add_para(doc, "API là thành phần trung gian giúp tầng giao diện ReactJS giao tiếp với tầng xử lý nghiệp vụ Node.js/Express.js. Thông qua API, frontend có thể gửi yêu cầu đến backend để thực hiện các chức năng như đăng nhập, tìm kiếm bác sĩ, đặt lịch khám, quản lý hồ sơ khám bệnh, nhận thông báo, upload tệp và quản trị dữ liệu hệ thống.")
add_para(doc, "Trong hệ thống “Ứng dụng đặt lịch khám bệnh cho phòng khám nhỏ”, API được thiết kế theo hướng RESTful, sử dụng giao thức HTTP và trao đổi dữ liệu dưới dạng JSON. Các API được tổ chức thành nhiều nhóm chức năng khác nhau, tương ứng với các nghiệp vụ chính của hệ thống. Việc phân nhóm API giúp mã nguồn backend rõ ràng hơn, thuận tiện cho quá trình phát triển, kiểm thử và bảo trì hệ thống.")
add_para(doc, "Các API quan trọng được bảo vệ bằng cơ chế xác thực JWT. Sau khi đăng nhập thành công, người dùng nhận được token và sử dụng token này trong các request tiếp theo để truy cập các chức năng cần đăng nhập. Ngoài ra, hệ thống còn áp dụng phân quyền theo vai trò người dùng gồm bệnh nhân, bác sĩ và quản trị viên nhằm bảo đảm mỗi nhóm người dùng chỉ được sử dụng các chức năng phù hợp với quyền hạn của mình.")

add_heading(doc, "4.3.1. Nguyên tắc thiết kế API")
add_para(doc, "Các API trong hệ thống được xây dựng dựa trên một số nguyên tắc chính nhằm bảo đảm tính thống nhất và dễ mở rộng. Trước hết, hệ thống sử dụng mô hình RESTful API, trong đó mỗi endpoint đại diện cho một tài nguyên hoặc một nhóm nghiệp vụ cụ thể. Các phương thức HTTP như GET, POST, PUT, PATCH và DELETE được sử dụng phù hợp với từng loại thao tác.")
add_para(doc, "Dữ liệu trao đổi giữa frontend và backend được định dạng dưới dạng JSON. Với các API cần truyền tệp như ảnh đại diện, ảnh cơ sở khám, ảnh bài viết hoặc tệp đính kèm hồ sơ khám bệnh, hệ thống sử dụng multipart/form-data. Các request gửi lên backend đều được kiểm tra dữ liệu đầu vào thông qua middleware validate trước khi chuyển đến controller xử lý nghiệp vụ.")
add_table(
    doc,
    ["Nội dung", "Mô tả"],
    [
        ("Kiểu thiết kế", "RESTful API"),
        ("Định dạng dữ liệu", "JSON, multipart/form-data đối với upload tệp"),
        ("Giao thức", "HTTP/HTTPS"),
        ("Backend", "Node.js, Express.js"),
        ("Cơ sở dữ liệu", "MongoDB thông qua Mongoose"),
        ("Xác thực", "JWT token"),
        ("Phân quyền", "Theo vai trò patient, doctor, admin"),
        ("Kiểm tra dữ liệu", "Sử dụng middleware validate trước khi xử lý request"),
        ("Phản hồi", "Trả về dữ liệu và thông báo xử lý cho frontend"),
    ],
    [Cm(4.5), Cm(12.5)],
)

add_heading(doc, "4.3.2. Nhóm API xác thực và phân quyền")
add_para(doc, "Nhóm API xác thực và phân quyền chịu trách nhiệm quản lý quá trình đăng ký, đăng nhập, xác thực email, quên mật khẩu, đổi mật khẩu và lấy thông tin người dùng hiện tại. Đây là nhóm API nền tảng vì hầu hết các chức năng quan trọng của hệ thống đều yêu cầu người dùng đăng nhập trước khi sử dụng.")
add_para(doc, "Sau khi đăng nhập thành công, backend trả về token xác thực cho frontend. Token này được gửi kèm trong các request tiếp theo để hệ thống xác định danh tính và vai trò của người dùng. Dựa trên vai trò đó, backend kiểm tra người dùng có quyền truy cập chức năng tương ứng hay không.")
add_table(
    doc,
    ["Chức năng", "Method", "Endpoint", "Vai trò sử dụng", "Mô tả"],
    [
        ("Đăng ký tài khoản", "POST", "/api/auth/register", "Khách truy cập", "Tạo tài khoản người dùng mới"),
        ("Đăng nhập", "POST", "/api/auth/login", "Tất cả người dùng", "Xác thực tài khoản và trả về token"),
        ("Xác thực email", "POST", "/api/auth/verify-email", "Người dùng đăng ký", "Kiểm tra mã OTP xác thực email"),
        ("Gửi lại OTP xác thực", "POST", "/api/auth/resend-verification-otp", "Người dùng đăng ký", "Gửi lại mã xác thực email"),
        ("Quên mật khẩu", "POST", "/api/auth/forgot-password", "Người dùng", "Gửi mã OTP đặt lại mật khẩu"),
        ("Đặt lại mật khẩu", "POST", "/api/auth/reset-password", "Người dùng", "Cập nhật mật khẩu mới sau khi xác thực OTP"),
        ("Đổi mật khẩu ban đầu", "POST", "/api/auth/change-initial-password", "Người dùng đã đăng nhập", "Đổi mật khẩu tạm thời do quản trị viên tạo"),
        ("Đổi mật khẩu", "PATCH", "/api/auth/change-password", "Người dùng đã đăng nhập", "Cập nhật mật khẩu hiện tại"),
        ("Lấy thông tin người dùng", "GET", "/api/auth/me", "Người dùng đã đăng nhập", "Trả về thông tin tài khoản hiện tại"),
    ],
    [Cm(3.6), Cm(1.8), Cm(5.4), Cm(3.4), Cm(4.2)],
)

add_heading(doc, "4.3.3. Nhóm API dữ liệu nền tảng")
add_para(doc, "Nhóm API dữ liệu nền tảng cung cấp các chức năng truy xuất và quản lý thông tin cơ bản của hệ thống như cơ sở khám, chuyên khoa, bác sĩ, gói khám và tài khoản người dùng. Một phần API trong nhóm này được mở công khai để khách truy cập và bệnh nhân có thể tìm kiếm thông tin khám bệnh. Các API tạo, cập nhật hoặc xóa dữ liệu thường được giới hạn cho quản trị viên.")
add_table(
    doc,
    ["Chức năng", "Method", "Endpoint", "Vai trò sử dụng", "Mô tả"],
    [
        ("Lấy danh sách cơ sở khám", "GET", "/api/clinics", "Công khai", "Hiển thị danh sách cơ sở khám"),
        ("Xem chi tiết cơ sở khám", "GET", "/api/clinics/:id", "Công khai", "Lấy thông tin chi tiết một cơ sở khám"),
        ("Tạo/cập nhật/xóa cơ sở khám", "POST/PUT/DELETE", "/api/clinics", "Admin", "Quản lý dữ liệu cơ sở khám"),
        ("Lấy danh sách chuyên khoa", "GET", "/api/specialties", "Công khai", "Hiển thị danh sách chuyên khoa"),
        ("Tạo chuyên khoa", "POST", "/api/specialties", "Admin", "Thêm chuyên khoa mới"),
        ("Lấy danh sách bác sĩ", "GET", "/api/doctors", "Công khai", "Hiển thị và tìm kiếm bác sĩ"),
        ("Xem chi tiết bác sĩ", "GET", "/api/doctors/:id", "Công khai", "Lấy thông tin chi tiết bác sĩ"),
        ("Quản lý bác sĩ", "POST/PUT/DELETE", "/api/doctors", "Admin", "Thêm, cập nhật, xóa hồ sơ bác sĩ"),
        ("Lấy danh sách gói khám", "GET", "/api/service-packages", "Công khai", "Hiển thị các gói khám đang hoạt động"),
        ("Xem chi tiết gói khám", "GET", "/api/service-packages/:id", "Công khai", "Lấy thông tin chi tiết gói khám"),
        ("Lấy thông tin cá nhân", "GET", "/api/users/me", "Người dùng đã đăng nhập", "Trả về thông tin cá nhân"),
        ("Cập nhật thông tin cá nhân", "PUT", "/api/users/me", "Người dùng đã đăng nhập", "Cập nhật hồ sơ cá nhân"),
    ],
    [Cm(4.0), Cm(2.0), Cm(5.4), Cm(3.2), Cm(4.6)],
)

add_heading(doc, "4.3.4. Nhóm API đặt lịch, hàng đợi và hồ sơ khám bệnh")
add_para(doc, "Đây là nhóm API nghiệp vụ chính của hệ thống. Nhóm API này xử lý các chức năng liên quan đến đặt lịch khám, quản lý trạng thái lịch hẹn, xử lý hàng đợi khám và lập hồ sơ khám bệnh sau khi bác sĩ hoàn tất lượt khám.")
add_para(doc, "Đối với bệnh nhân, nhóm API này hỗ trợ tạo lịch hẹn, xem lịch hẹn cá nhân, yêu cầu hủy hoặc đổi lịch, tham gia danh sách chờ và xem hồ sơ khám bệnh. Đối với bác sĩ và quản trị viên, nhóm API này hỗ trợ theo dõi lịch khám, cập nhật trạng thái khám, tạo hồ sơ khám bệnh, cập nhật hồ sơ và xuất phiếu kết quả khám.")
add_table(
    doc,
    ["Chức năng", "Method", "Endpoint", "Vai trò sử dụng", "Mô tả"],
    [
        ("Tạo lịch hẹn", "POST", "/api/appointments", "Patient", "Bệnh nhân đặt lịch khám"),
        ("Xem lịch hẹn của tôi", "GET", "/api/appointments/my", "Patient", "Lấy danh sách lịch hẹn của bệnh nhân"),
        ("Xem danh sách lịch hẹn", "GET", "/api/appointments", "Admin", "Quản trị viên xem danh sách lịch hẹn"),
        ("Xem lịch hôm nay", "GET", "/api/appointments/clinic/:clinicId/today", "Doctor, Admin", "Lấy danh sách lịch hẹn trong ngày theo cơ sở"),
        ("Hủy lịch hẹn", "PATCH", "/api/appointments/:id/cancel", "Patient", "Gửi yêu cầu hủy lịch hẹn"),
        ("Yêu cầu đổi lịch", "PATCH", "/api/appointments/:id/reschedule-request", "Patient", "Gửi yêu cầu thay đổi ngày hoặc giờ khám"),
        ("Cập nhật trạng thái khám", "PATCH", "/api/appointments/:id/consultation-status", "Doctor, Admin", "Cập nhật trạng thái tiếp nhận và khám"),
        ("Cập nhật trạng thái lịch hẹn", "PATCH", "/api/appointments/:id/status", "Doctor, Admin", "Xác nhận, hoàn tất hoặc thay đổi trạng thái lịch"),
        ("Tạo lượt chờ khám", "POST", "/api/waiting-list", "Patient", "Đăng ký vào danh sách chờ"),
        ("Xem danh sách chờ của tôi", "GET", "/api/waiting-list/my", "Patient", "Lấy danh sách chờ của bệnh nhân"),
        ("Chấp nhận lịch từ hàng chờ", "POST", "/api/waiting-list/:id/accept", "Patient", "Chấp nhận khung giờ được đề xuất"),
        ("Từ chối lịch từ hàng chờ", "POST", "/api/waiting-list/:id/decline", "Patient", "Từ chối khung giờ được đề xuất"),
        ("Hủy lượt chờ", "DELETE", "/api/waiting-list/:id", "Patient", "Hủy lượt chờ khám"),
        ("Tạo hồ sơ khám bệnh", "POST", "/api/medical-records", "Doctor", "Bác sĩ lập hồ sơ khám bệnh"),
        ("Xem hồ sơ khám của tôi", "GET", "/api/medical-records/my", "Patient", "Bệnh nhân xem hồ sơ khám bệnh"),
        ("Xem hồ sơ tái khám", "GET", "/api/medical-records/follow-ups/my", "Patient", "Lấy danh sách hồ sơ cần theo dõi tái khám"),
        ("Cập nhật hồ sơ khám bệnh", "PATCH", "/api/medical-records/:id", "Doctor", "Bác sĩ cập nhật hồ sơ khám"),
        ("Cập nhật tệp đính kèm", "PATCH", "/api/medical-records/:id/attachments", "Patient, Doctor, Admin", "Cập nhật tệp đính kèm hồ sơ khám"),
        ("Xem chi tiết hồ sơ", "GET", "/api/medical-records/:id", "Patient, Doctor, Admin", "Lấy chi tiết hồ sơ khám bệnh"),
    ],
    [Cm(3.8), Cm(1.8), Cm(5.6), Cm(3.2), Cm(4.4)],
)

add_heading(doc, "4.3.5. Nhóm API thông báo, upload, PDF và AI")
add_para(doc, "Nhóm API này hỗ trợ các chức năng bổ sung nhằm hoàn thiện trải nghiệm sử dụng hệ thống. Các API thông báo giúp người dùng nhận được thông tin về lịch hẹn, hàng chờ, yêu cầu hủy lịch, đổi lịch hoặc nhắc tái khám. Các API upload dùng để tải ảnh và tệp đính kèm lên hệ thống. Bên cạnh đó, hệ thống còn hỗ trợ xuất PDF hồ sơ khám bệnh và tích hợp AI để tư vấn triệu chứng ban đầu cho người dùng.")
add_table(
    doc,
    ["Chức năng", "Method", "Endpoint", "Vai trò sử dụng", "Mô tả"],
    [
        ("Xem thông báo", "GET", "/api/notifications", "Patient, Doctor, Admin", "Lấy danh sách thông báo của người dùng"),
        ("Xem thông báo của tôi", "GET", "/api/notifications/my", "Patient, Doctor, Admin", "Lấy danh sách thông báo cá nhân"),
        ("Đánh dấu đã đọc tất cả", "PATCH", "/api/notifications/read-all", "Patient, Doctor, Admin", "Đánh dấu toàn bộ thông báo là đã đọc"),
        ("Đánh dấu đã đọc một thông báo", "PATCH", "/api/notifications/:id/read", "Patient, Doctor, Admin", "Cập nhật trạng thái đã đọc"),
        ("Upload ảnh/tệp cơ sở khám, bác sĩ, chuyên khoa, gói khám, bài viết", "POST", "/api/uploads/...", "Admin, Doctor", "Tải các tệp hình ảnh lên hệ thống"),
        ("Upload tệp hồ sơ khám", "POST", "/api/uploads/medical-record-attachments", "Patient, Doctor, Admin", "Tải tệp đính kèm hồ sơ khám bệnh"),
        ("Upload ảnh đại diện người dùng", "POST", "/api/uploads/user-avatar", "Người dùng đã đăng nhập", "Tải ảnh đại diện cá nhân"),
        ("Xuất PDF lịch hẹn", "GET", "/api/appointments/:id/pdf", "Patient, Doctor, Admin", "Xuất thông tin lịch hẹn dạng PDF"),
        ("Xuất phiếu số thứ tự", "GET", "/api/appointments/:id/queue-ticket/pdf", "Patient, Doctor, Admin", "Xuất phiếu hàng đợi khám dạng PDF"),
        ("Xuất PDF hồ sơ khám", "GET", "/api/medical-records/:id/pdf", "Patient, Doctor, Admin", "Xuất phiếu kết quả khám bệnh"),
        ("Tư vấn triệu chứng", "POST", "/api/ai/symptom-checker", "Công khai / người dùng", "Phân tích triệu chứng ban đầu"),
        ("Trợ lý triệu chứng", "POST", "/api/ai/symptom-assistant", "Công khai / người dùng", "Gợi ý chuyên khoa hoặc hướng khám phù hợp"),
    ],
    [Cm(4.0), Cm(2.0), Cm(5.6), Cm(3.2), Cm(4.2)],
)

add_heading(doc, "4.3.6. Nhóm API quản trị hệ thống")
add_para(doc, "Nhóm API quản trị hệ thống dành cho quản trị viên, phục vụ việc theo dõi tình hình vận hành và quản lý dữ liệu trong toàn bộ hệ thống. Các API này cho phép quản trị viên xem thống kê tổng quan, quản lý tài khoản, bác sĩ, cơ sở khám, chuyên khoa, lịch làm việc, gói khám, bài viết, đánh giá và nhật ký thao tác hệ thống.")
add_table(
    doc,
    ["Chức năng", "Method", "Endpoint", "Vai trò sử dụng", "Mô tả"],
    [
        ("Xem tổng quan hệ thống", "GET", "/api/admin/dashboard", "Admin", "Lấy dữ liệu thống kê phục vụ trang quản trị"),
        ("Quản lý người dùng", "GET", "/api/admin/users", "Admin", "Lấy danh sách tài khoản người dùng"),
        ("Cập nhật trạng thái tài khoản", "PATCH", "/api/admin/users/:userId/status", "Admin", "Khóa hoặc kích hoạt tài khoản"),
        ("Xem danh sách tài khoản bác sĩ", "GET", "/api/admin/doctor-users", "Admin", "Xem các tài khoản gắn với bác sĩ"),
        ("Xem nhật ký hệ thống", "GET", "/api/admin/audit-logs", "Admin", "Lấy danh sách nhật ký thao tác"),
        ("Xem chi tiết nhật ký", "GET", "/api/admin/audit-logs/:id", "Admin", "Xem chi tiết một bản ghi nhật ký"),
        ("Tạo/cập nhật cơ sở khám", "POST/PATCH", "/api/admin/clinics", "Admin", "Quản lý dữ liệu cơ sở khám"),
        ("Tạo/cập nhật chuyên khoa", "POST/PATCH", "/api/admin/specialties", "Admin", "Quản lý dữ liệu chuyên khoa"),
        ("Tạo/cập nhật bác sĩ", "POST/PATCH", "/api/admin/doctors", "Admin", "Quản lý hồ sơ bác sĩ"),
        ("Tạo tài khoản bác sĩ", "POST", "/api/admin/doctors/:doctorId/account", "Admin", "Tạo tài khoản đăng nhập cho bác sĩ"),
        ("Liên kết tài khoản bác sĩ", "PATCH", "/api/admin/doctors/:doctorId/account", "Admin", "Liên kết hồ sơ bác sĩ với tài khoản người dùng"),
        ("Đặt lại mật khẩu bác sĩ", "PATCH", "/api/admin/doctors/:doctorId/account/reset-password", "Admin", "Tạo mật khẩu mới cho tài khoản bác sĩ"),
        ("Cập nhật trạng thái bác sĩ", "PATCH", "/api/admin/doctors/:doctorId/account/status", "Admin", "Kích hoạt hoặc vô hiệu hóa tài khoản bác sĩ"),
        ("Cập nhật lịch làm việc", "POST", "/api/admin/schedules", "Admin", "Thiết lập lịch làm việc cho bác sĩ"),
    ],
    [Cm(4.0), Cm(1.8), Cm(6.6), Cm(2.8), Cm(5.0)],
)

add_para(doc, "Nhìn chung, thiết kế API của hệ thống được tổ chức theo từng nhóm chức năng rõ ràng, phù hợp với kiến trúc frontend ReactJS và backend Node.js/Express.js. Việc sử dụng RESTful API, JWT và phân quyền theo vai trò giúp hệ thống dễ mở rộng, dễ bảo trì và bảo đảm an toàn trong quá trình xử lý dữ liệu.")

doc.save(OUT_PATH)
print(OUT_PATH)
