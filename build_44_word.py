from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\DELL\Documents\Codex\2026-05-27\clinic-booking")
SOURCE_MD = ROOT / "4.4_thiet_ke_luong_trinh_tu.md"
OUT_DOCX = ROOT / "4.4_thiet_ke_luong_trinh_tu.docx"


def set_font(run, name="Times New Roman", size=13, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(13)


def add_paragraph(doc, text="", align=None, size=13, bold=False, italic=False, space_after=6, space_before=0, line_spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        add_inline_text(p, text, size=size, bold=bold, italic=italic)
    return p


def add_inline_text(paragraph, text, size=13, bold=False, italic=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=max(10, size - 1), bold=bold, italic=italic)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, bold=bold, italic=italic)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 2 else 4)
    p.paragraph_format.line_spacing = 1.15
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=16, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_font(run, size=14, bold=True)
    else:
        run = p.add_run(text)
        set_font(run, size=13, bold=True)
    return p


def add_caption(doc, text):
    return add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        italic=True,
        space_after=4,
        line_spacing=1.15,
    )


def add_quote_note(doc, text):
    return add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        italic=True,
        space_before=2,
        space_after=2,
        line_spacing=1.15,
    )


def build_doc():
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    # Match the thesis-like source report structure.
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.27)
    sec.footer_distance = Cm(1.27)

    set_normal_style(doc)

    # Remove the default empty paragraph if present later by overwriting content flow.
    in_code = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.right_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.0
                for idx, code_line in enumerate(code_text.splitlines()):
                    if idx:
                        p.add_run().add_break()
                    run = p.add_run(code_line)
                    set_font(run, name="Consolas", size=10.5)
                continue
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph()
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            continue

        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
            continue

        if line.startswith("> "):
            note = line[2:].strip()
            note = note.replace("**", "").strip()
            add_quote_note(doc, note)
            continue

        if line.startswith("Hình "):
            add_caption(doc, line)
            continue

        add_paragraph(doc, line, line_spacing=1.5)

    # Clean leading/trailing empty paragraphs a little.
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
